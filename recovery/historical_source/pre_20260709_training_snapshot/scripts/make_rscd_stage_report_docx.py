from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


OUT = Path(r"C:\Users\DELL\Desktop\RSCD视觉道路摩擦可供性估计阶段性进展.docx")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text


def main() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "SimSun"
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("RSCD视觉道路摩擦可供性估计阶段性进展", level=1)
    doc.add_paragraph(
        "本阶段的核心目标是把原有视觉摩擦可供性算法放到RSCD公开数据集的标准分类协议下重新审视，"
        "明确哪些结果可以和公开论文公平比较，哪些结果只能作为弱摩擦区间任务的内部指标。"
    )

    doc.add_heading("一、数据集与协议判断", level=2)
    add_table(
        doc,
        ["数据集", "本地规模", "图像形态", "当前用途"],
        [
            ["RSCD balanced", "5400 / 5400 / 5400", "240x360道路表面patch", "快速摩擦可供性与消融实验"],
            ["RSCD full", "958941 / 19860 / 49500", "240x360道路表面patch", "RSCD-27公开对标协议"],
            ["RSCD per-day", "678281 / 290227 / 59793", "240x360道路表面patch", "更严格的按采集日期泛化协议"],
            ["RoadSaW", "2400 / 2400 / 1728", "400x400 BEV道路湿滑patch", "湿滑路面压力测试"],
            ["RoadSC", "600 / 600 / 573", "900x900雪地道路patch", "冬季低附着压力测试"],
        ],
    )
    doc.add_paragraph(
        "结论：RSCD不是完整左右轮前方视野，而是道路表面裁剪patch；RoadSaW和RoadSC更接近方形BEV道路patch。"
        "三者存在明显视角、尺寸、标签体系和成像风格差异，因此不能把多数据集混合作为唯一主结论。"
        "合理路线是：RSCD做公开分类对标，RoadSaW/RoadSC做湿滑和雪地泛化压力测试。"
    )

    doc.add_heading("二、RSCD论文常用评价指标与SOTA边界", level=2)
    add_table(
        doc,
        ["指标", "含义", "备注"],
        [
            ["Top-1 Accuracy", "最高概率类别是否等于真实类别", "RSCD论文最常见主指标"],
            ["Mean Precision", "逐类precision后取平均", "关注类别均衡性"],
            ["Mean Recall", "逐类recall后取平均", "关注少数类召回"],
            ["Mean F1 / Macro-F1", "逐类F1后取平均", "比Top-1更适合细粒度不均衡任务"],
            ["Confusion Matrix", "观察类别混淆", "用于解释湿滑、积水、材料混淆"],
        ],
    )
    doc.add_paragraph(
        "公开RSCD文献中，RoadFormer报告RSCD Top-1约92.52%；RoadMamba报告更强的RSCD结果，"
        "其中RoadMamba-B Top-1约92.81%，RoadMamba-T消融行约Top-1 91.52%、Mean-F1 82.44%。"
        "这些数字只能在相同split、相同27类标签、相同训练设置和相同指标下公平比较。"
    )

    doc.add_heading("三、当前算法设计", level=2)
    doc.add_paragraph("基础模型为ConvNeXt-Tiny全局视觉特征分类器：")
    doc.add_paragraph("图像 x -> ConvNeXt -> LayerNorm -> Dropout -> 27类分类头")
    doc.add_paragraph("改进模型加入PhysicsTexture分支：")
    doc.add_paragraph("图像 x -> ConvNeXt全局特征 f_g；同时 x -> PhysicsTexture物理纹理特征 f_p；拼接 [f_g, f_p] 后分类。")
    doc.add_paragraph(
        "PhysicsTexture不是普通CNN分支，而是固定可微的道路物理纹理描述，包括灰度、饱和度、亮度、Sobel梯度、"
        "Laplacian、雪白低饱和区域、镜面高光、暗水区域、湿滑连通性、雪区域连通性、低纹理、近白过曝、"
        "薄水膜和上下区域差异等。它的目标是让模型直接看到与干湿、积水、雪、冰、材料粗糙度相关的证据。"
    )

    doc.add_heading("四、快速验证结果", level=2)
    add_table(
        doc,
        ["模型", "Top-1", "Macro-F1", "Weighted-F1", "Balanced Acc", "决策"],
        [
            ["ConvNeXt-Tiny", "70.37%", "70.15%", "70.15%", "70.37%", "baseline"],
            ["+ PhysicsTexture + quality cues", "72.54%", "72.62%", "72.62%", "72.54%", "晋级正式实验"],
            ["+ Factor aux + Physics aux", "71.72%", "71.72%", "71.72%", "71.72%", "当前形式舍弃"],
        ],
    )
    doc.add_paragraph(
        "PhysicsTexture相对ConvNeXt-Tiny提升Top-1 2.17个百分点、Macro-F1 2.47个百分点。"
        "类别分析显示，提升主要集中在water_mud、wet_asphalt_severe、water_concrete、wet_gravel等湿滑、积水和纹理混合类别，"
        "这与模块设计目标一致。Factor auxiliary虽然有理论动机，但当前简单辅助CE形式没有超过纯PhysicsTexture，因此暂不作为正式主模型。"
    )

    doc.add_heading("五、当前正式实验状态", level=2)
    doc.add_paragraph(
        "已启动RSCD-27正式pair实验：formal ConvNeXt-Tiny baseline和formal ConvNeXt-Tiny + PhysicsTexture + quality cues。"
        "二者使用相同full manifest、相同image size=192、相同samples_per_epoch=36000、相同batch=12、accum=2、相同epoch预算和相同指标。"
    )
    add_table(
        doc,
        ["模型", "Epoch", "Val Top-1", "Val Macro-F1", "Val Balanced Acc", "状态"],
        [
            ["ConvNeXt-Tiny", "1", "73.82%", "72.08%", "72.83%", "已完成"],
            ["ConvNeXt-Tiny", "2", "77.89%", "76.23%", "77.18%", "已完成"],
            ["ConvNeXt-Tiny", "3", "79.06%", "77.68%", "78.65%", "已完成"],
            ["ConvNeXt-Tiny", "4", "79.80%", "78.41%", "79.23%", "已完成"],
            ["ConvNeXt-Tiny", "5", "81.86%", "80.29%", "81.58%", "已完成"],
            ["ConvNeXt-Tiny", "6", "82.40%", "80.93%", "82.03%", "已完成"],
            ["ConvNeXt-Tiny + PhysicsTexture", "1", "75.09%", "73.04%", "73.70%", "已完成"],
            ["ConvNeXt-Tiny + PhysicsTexture", "2", "77.53%", "75.57%", "75.79%", "已完成"],
            ["ConvNeXt-Tiny + PhysicsTexture", "3", "79.83%", "78.37%", "78.48%", "已完成"],
            ["ConvNeXt-Tiny + PhysicsTexture", "4", "80.57%", "79.17%", "79.82%", "已完成"],
            ["ConvNeXt-Tiny + PhysicsTexture", "5", "82.30%", "80.95%", "81.53%", "已完成"],
        ],
    )
    doc.add_paragraph(
        "正式早期结果目前是有希望但仍需等待最终test的证据：第1轮完整验证集上，PhysicsTexture相对baseline提升Top-1 1.27个百分点、"
        "Macro-F1 0.96个百分点；第2轮略低于baseline；第3轮又超过同epoch baseline，Top-1提升0.77个百分点、"
        "Macro-F1提升0.69个百分点；第4轮继续超过同epoch baseline，Macro-F1提升0.76个百分点。"
        "第5轮PhysicsTexture继续超过同epoch baseline，Macro-F1提升0.66个百分点；而且PhysicsTexture第5轮Macro-F1 80.95%已经略高于baseline第6轮80.93%。"
        "因此可以说PhysicsTexture在正式验证集上已经形成Macro-F1优势趋势，但最终结论仍必须等待两个模型的best checkpoint和test结果。"
    )
    doc.add_paragraph(
        "正式实验输出目录：D:\\NMI_SPWFM_datasets\\friction_affordance_outputs\\rscd_surface_classification。"
        "日志目录：E:\\perception\\friction_affordance_field\\outputs\\rscd_surface_formal_queue。"
    )

    doc.add_heading("六、下一步计划", level=2)
    for item in [
        "等待formal pair完成，得到正式Top-1、Macro-F1、Weighted-F1和Balanced Accuracy。",
        "如果PhysicsTexture正式结果复现fast提升，则作为RSCD-27公开对标表的核心方法。",
        "进一步做类级别混淆分析，重点解释wet/water/snow/ice/material相关类别的提升和下降。",
        "若正式结果不足，下一步尝试更强但仍可解释的结构：因子乘积分解分类器、条件材料-湿滑解码器、局部区域token池化。",
        "RoadSaW/RoadSC不与RSCD混成一个主指标，而作为湿滑/雪地跨数据集压力测试和泛化失败分析。",
    ]:
        doc.add_paragraph(item, style=None)

    doc.add_heading("七、当前结论", level=2)
    doc.add_paragraph(
        "当前最可信结论是：PhysicsTexture模块不仅在弱摩擦可供性任务中有效，在RSCD原始27类公开分类协议的快速公平对比中也带来稳定提升。"
        "这为后续论文提供了一个清晰故事：道路摩擦相关视觉识别不能只依赖全局语义特征，还需要显式建模湿滑、积水、雪冰、粗糙度、纹理和反射等物理纹理证据。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
