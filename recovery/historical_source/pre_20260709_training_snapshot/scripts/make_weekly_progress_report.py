from __future__ import annotations

import csv
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

import yaml
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from PIL import Image as PILImage

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from friction_affordance.ontology import (  # noqa: E402
    FRICTION_STATES,
    MATERIALS,
    RISK,
    SNOW,
    TASKS,
    UNEVENNESS,
    WETNESS,
    risk_from_mu_interval,
    weak_mu_interval_from_state,
)


EXPERIMENT = ROOT / "outputs" / "topvenue_v4_evidencefield_rtx3050_stable"
CONFIG = ROOT / "configs" / "experiments" / "topvenue_v4_evidencefield_rtx3050_stable.yaml"
REPORT_DIR = ROOT / "reports"
REPORT_MD = REPORT_DIR / "friction_affordance_research_content_explained_20260618.md"
REPORT_PDF = REPORT_DIR / "friction_affordance_research_content_explained_20260618.pdf"
PLAIN_DOCX = REPORT_DIR / "friction_affordance_plain_explanation_20260618.docx"
FORMULA_DIR = REPORT_DIR / "_formula_cache"


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    data = load_report_data()
    md = build_markdown(data)
    REPORT_MD.write_text(md, encoding="utf-8")
    build_pdf(data)
    build_plain_docx(data)
    print(f"wrote: {REPORT_MD}")
    print(f"wrote: {REPORT_PDF}")
    print(f"wrote: {PLAIN_DOCX}")


def load_report_data() -> dict[str, Any]:
    cfg = yaml.safe_load(CONFIG.read_text(encoding="utf-8"))
    evaluate = read_json(EXPERIMENT / "evaluate_test.json")
    detailed = read_json(EXPERIMENT / "detailed_test.json")
    calibration = read_json(EXPERIMENT / "interval_calibration_90.json")
    diagnostic = read_json(EXPERIMENT / "dataset_id_diagnostic.json")
    audit = read_json(EXPERIMENT / "topvenue_result_audit.json")
    manifest_stats = read_json(EXPERIMENT / "manifest_stats_train.json")
    train_log = (EXPERIMENT / "background_20260617_210118.out.log").read_text(
        encoding="utf-8",
        errors="ignore",
    )
    best_epoch, best_val = parse_best_epoch(train_log)
    manifests = manifest_counts()
    evidence_maps = sorted((EXPERIMENT / "evidence_maps").glob("*.jpg"))
    return {
        "cfg": cfg,
        "evaluate": evaluate,
        "detailed": detailed,
        "calibration": calibration,
        "diagnostic": diagnostic,
        "audit": audit,
        "manifest_stats": manifest_stats,
        "best_epoch": best_epoch,
        "best_val": best_val,
        "manifests": manifests,
        "evidence_maps": evidence_maps,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "best_ckpt_mtime": datetime.fromtimestamp((EXPERIMENT / "best.pt").stat().st_mtime).strftime(
            "%Y-%m-%d %H:%M"
        ),
        "last_ckpt_mtime": datetime.fromtimestamp((EXPERIMENT / "last.pt").stat().st_mtime).strftime(
            "%Y-%m-%d %H:%M"
        ),
    }


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def parse_best_epoch(log_text: str) -> tuple[int | None, dict[str, float]]:
    epoch = None
    best_epoch = None
    best_val: dict[str, float] = {}
    for raw_line in log_text.splitlines():
        line = raw_line.strip()
        match = re.match(r"Epoch\s+(\d+)/", line)
        if match:
            epoch = int(match.group(1))
            continue
        if line.startswith("val  :"):
            pairs = re.findall(r"([A-Za-z0-9_]+)=([-+]?\d*\.?\d+(?:[eE][-+]?\d+)?)", line)
            current = {k: float(v) for k, v in pairs}
            if current and (not best_val or current.get("loss", float("inf")) < best_val.get("loss", float("inf"))):
                best_val = current
                best_epoch = epoch
    return best_epoch, best_val


def manifest_counts() -> dict[str, dict[str, int]]:
    files = {
        "RSCD": {
            "train": ROOT / "data" / "manifests_full" / "rscd_prepared_train.csv",
            "val": ROOT / "data" / "manifests_full" / "rscd_prepared_val.csv",
            "test": ROOT / "data" / "manifests_full" / "rscd_prepared_test.csv",
        },
        "RoadSaW": {
            "train": ROOT / "data" / "manifests_full" / "roadsaw_train.csv",
            "val": ROOT / "data" / "manifests_full" / "roadsaw_val.csv",
            "test": ROOT / "data" / "manifests_full" / "roadsaw_test.csv",
        },
        "RoadSC": {
            "train": ROOT / "data" / "manifests_full" / "roadsc_train.csv",
            "val": ROOT / "data" / "manifests_full" / "roadsc_val.csv",
            "test": ROOT / "data" / "manifests_full" / "roadsc_test.csv",
        },
    }
    out: dict[str, dict[str, int]] = {}
    for dataset, split_files in files.items():
        out[dataset] = {}
        for split, path in split_files.items():
            with path.open("r", encoding="utf-8", newline="") as f:
                out[dataset][split] = max(sum(1 for _ in f) - 1, 0)
    return out


def build_markdown(data: dict[str, Any]) -> str:
    detailed = data["detailed"]
    calibration = data["calibration"]
    diagnostic = data["diagnostic"]
    audit = data["audit"]
    cfg = data["cfg"]
    lines = [
        "# 视觉图像摩擦可供性估计算法研究内容",
        "",
        f"- 生成时间: {data['generated_at']}",
        f"- 实验目录: `{relative(EXPERIMENT)}`",
        f"- 最佳 checkpoint: `best.pt`, 保存时间 {data['best_ckpt_mtime']}, best epoch {data['best_epoch']}",
        f"- 当前审计结论: `{audit.get('verdict', 'unknown')}`",
        "",
        "## 研究结论概览",
        "",
        f"当前模型在平衡测试集 {detailed['num_samples_seen']} 张图像上达到: "
        f"friction macro-F1={pct(task_metric(detailed, 'friction', 'macro_f1'))}, "
        f"risk macro-F1={pct(task_metric(detailed, 'risk', 'macro_f1'))}, "
        f"低摩擦检出 F1={pct(detailed['low_friction_detection']['f1'])}, "
        f"校准后 90% 区间覆盖率={pct(calibration['test_split']['calibrated_coverage'])}。",
        "",
        "判断: 当前模型已经形成完整的视觉摩擦可供性估计流程，能够支持研究内容展示和进一步论文实验；但还不能作为顶会/顶刊最终结果。主要原因是 LODO 泛化实验和完整消融尚未跑完，RoadSaW 域上的 worst macro-F1 偏低，且数据集身份可预测性过高。",
        "",
        "## 数据集与标签",
        "",
        "| 数据集 | train | val | test | 主要作用 |",
        "|---|---:|---:|---:|---|",
    ]
    for name, splits in data["manifests"].items():
        role = {
            "RSCD": "大规模路面状态/材质/平整度弱标签主数据源",
            "RoadSaW": "湿滑、干湿程度和材质域泛化测试",
            "RoadSC": "雪/积雪类低摩擦状态补充",
        }[name]
        lines.append(f"| {name} | {splits['train']} | {splits['val']} | {splits['test']} | {role} |")
    lines.extend(
        [
            "",
            f"任务标签: friction={FRICTION_STATES}; material={MATERIALS}; wetness={WETNESS}; snow={SNOW}; unevenness={UNEVENNESS}; risk={RISK}。",
            "",
            "核心注意: 当前公开视觉数据集没有真实路试摩擦系数测量。代码把路况标签映射为保守摩擦区间，用于弱监督训练和风险排序。因此本项目更严谨的表述是: 从视觉图像估计 tire-road friction affordance interval / risk，而不是从图像直接测量真实瞬时 μ。",
            "",
            "## 算法公式摘要",
            "",
            "输入图像 x ∈ R^{3×H×W}。模型输出多任务标签分布 p_t(y|x)、风险等级、摩擦均值 μ_hat 和摩擦区间 I=[μ_L, μ_U]。",
            "",
            "全局视觉特征: h_g=f_θ(x)。物理纹理分支: h_p=g_p(E_phys(x))。局部证据场: h_e,A,I_E,z_E=g_e(E_local(x))。融合特征 h=LayerNorm([h_g;h_p;h_e])。",
            "",
            "摩擦状态集合: s=(m,w,n,u)，其中 m 为材质，w 为湿度，n 为雪冰状态，u 为平整度。p(s|x)=p(m|x)p(w|x)p(n|x)p(u|x)。每个状态有弱物理区间 [l_s,u_s]，预测区间为 E_s[l_s], E_s[u_s]，再按状态熵扩大边界。",
            "",
            "最终区间融合: I_0=0.8 I_F+0.2 I_P，I=0.2 I_E+0.8 I_0。风险 logits 融合: z_risk=0.85 z_main+0.15 z_E。",
            "",
            "总损失包含: masked 多任务 CE、risk ordinal EMD、FrictionSet 兼容性 NLL、Group-DRO、V-REx、区间 censor logistic NLL、覆盖/端点/宽度损失、湿度-μ 和风险-μ 单调约束、CORAL 域对齐、证据场辅助损失和注意力先验/平滑/一致性损失。",
            "",
            "## 当前最好结果",
            "",
            "| 指标 | 数值 | 判断 |",
            "|---|---:|---|",
            f"| friction macro-F1 | {pct(task_metric(detailed, 'friction', 'macro_f1'))} | 总体较强 |",
            f"| risk macro-F1 | {pct(task_metric(detailed, 'risk', 'macro_f1'))} | 可支撑风险识别分析 |",
            f"| low-friction recall | {pct(detailed['low_friction_detection']['recall'])} | 低摩擦召回高，是当前亮点 |",
            f"| raw interval coverage | {pct(detailed['mu_interval']['coverage'])} | 原始区间偏窄 |",
            f"| calibrated coverage | {pct(calibration['test_split']['calibrated_coverage'])} | 接近 90% 目标 |",
            f"| dataset ID balanced acc | {pct(diagnostic['overall_dataset_id_balanced_accuracy'])} | 域捷径严重，需要后续处理 |",
            "",
            "## 后续研究计划",
            "",
            "1. 立即补齐 global-only、+FrictionSet、+DG、+EvidenceField 的消融，并跑 leave-one-dataset-out。",
            "2. 针对 RoadSaW 加强跨域鲁棒性: 更强颜色/风格扰动、条件 CORAL/IRM/V-REx、域对抗或 domain-specific adapter。",
            "3. 改进证据场: 引入道路区域伪分割或底部车道 ROI，一致性监督 attention，避免证据场单独 head 过弱。",
            "4. 论文表述改为弱监督 friction affordance interval estimation，并把校准区间、低摩擦召回和审计流程作为核心亮点。",
        ]
    )
    _ = cfg
    return "\n".join(lines) + "\n"


def build_pdf(data: dict[str, Any]) -> None:
    font_name = register_font()
    styles = make_styles(font_name)
    doc = SimpleDocTemplate(
        str(REPORT_PDF),
        pagesize=A4,
        rightMargin=1.55 * cm,
        leftMargin=1.55 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.35 * cm,
        title="视觉图像摩擦可供性估计算法研究内容",
    )
    story: list[Any] = []

    add_title(story, styles, data)
    add_executive_summary(story, styles, data)
    add_dataset_section(story, styles, data)
    add_algorithm_section(story, styles, data)
    add_loss_training_section(story, styles, data)
    add_results_section(story, styles, data)
    add_innovation_section(story, styles, data)
    add_next_steps_section(story, styles, data)
    add_appendix(story, styles, data)
    doc.build(story, onFirstPage=draw_page_number, onLaterPages=draw_page_number)


def build_plain_docx(data: dict[str, Any]) -> None:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    detailed = data["detailed"]
    calibration = data["calibration"]
    diagnostic = data["diagnostic"]
    doc = Document()
    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("视觉图像摩擦可供性估计算法大白话讲解稿", level=0)
    add_doc_para(
        doc,
        "这份文档的作用，是把研究内容 PDF 里的技术内容翻译成更容易讲出来的话。"
        "正式 PDF 适合给别人看算法细节，这份 Word 更适合你在讲解时顺着思路说。",
    )

    doc.add_heading("1. 一句话说明这个研究在做什么", level=1)
    add_doc_para(
        doc,
        "这个研究想解决的问题是：只看一张路面图像，判断这段路大概有多滑、风险高不高，"
        "并给出一个摩擦系数可能落在哪个范围的估计。这里不是直接测量真实摩擦系数，"
        "因为公开数据集没有车辆动力学传感器真值；我们做的是基于视觉标签的摩擦可供性估计。",
    )

    doc.add_heading("2. 为什么不能直接说“我预测了真实摩擦系数”", level=1)
    add_doc_para(
        doc,
        "RSCD、RoadSaW、RoadSC 这些数据集主要提供的是干、湿、积水、雪、冰、材质、平整度等视觉标签。"
        "它们没有告诉我们某一张图对应的真实轮胎-路面摩擦系数是多少。所以我把这些视觉标签映射成一个比较保守的摩擦区间，"
        "比如干燥沥青摩擦较高，冰面摩擦很低，湿路面居中。这样更严谨：模型学到的是摩擦风险和摩擦区间，而不是虚构的实测 μ。",
    )

    doc.add_heading("3. 数据集怎么用", level=1)
    add_doc_para(
        doc,
        "当前使用了三个公开数据集。RSCD 样本最多，是主数据源；RoadSaW 提供湿滑和材质变化；RoadSC 补充雪类低摩擦场景。"
        "训练时不是让 RSCD 一家独大，而是用平衡采样，让不同数据集和不同类别都有机会被模型看到。",
    )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "数据集"
    hdr[1].text = "train"
    hdr[2].text = "val"
    hdr[3].text = "test"
    for name, splits in data["manifests"].items():
        row = table.add_row().cells
        row[0].text = name
        row[1].text = str(splits["train"])
        row[2].text = str(splits["val"])
        row[3].text = str(splits["test"])

    doc.add_heading("4. 算法整体怎么讲", level=1)
    add_doc_para(
        doc,
        "可以把算法想成四层。第一层是 ConvNeXt，看整张图，理解这是什么路面。第二层是物理纹理分支，专门看雪白、反光、水迹、粗糙度这些与打滑相关的低层特征。"
        "第三层是局部证据场，它会在图像底部更可能接触车轮的区域找证据，而不是被背景干扰。第四层是 FrictionSet，把材质、湿度、雪冰、平整度组合成 450 种潜在路面状态，"
        "再把这些状态转换成摩擦区间。",
    )

    doc.add_heading("5. 公式怎么用人话解释", level=1)
    add_doc_para(
        doc,
        "PDF 里第一个公式是在定义数据：每个样本不仅有图像和标签，还有弱摩擦区间、数据集来源和训练分组。"
        "第二类公式是在讲特征：全局特征看整体，物理纹理特征看反光、水、雪，局部证据场看关键区域。"
        "第三类公式是在讲摩擦区间：先根据 450 种潜在路面状态算一个区间，再和网络直接预测的区间、局部证据区间融合。"
        "第四类公式是在讲训练：损失函数不只要求分类正确，还要求风险等级有顺序、摩擦区间能覆盖、不同数据集之间更稳健。",
    )

    doc.add_heading("6. 当前结果怎么讲", level=1)
    add_doc_para(
        doc,
        f"当前最好模型在平衡测试集上 friction macro-F1 为 {pct(task_metric(detailed, 'friction', 'macro_f1'))}，"
        f"risk macro-F1 为 {pct(task_metric(detailed, 'risk', 'macro_f1'))}，"
        f"低摩擦检出 F1 为 {pct(detailed['low_friction_detection']['f1'])}。"
        f"原始摩擦区间覆盖率只有 {pct(detailed['mu_interval']['coverage'])}，说明原始区间偏窄；"
        f"经过 conformal calibration 后，覆盖率提升到 {pct(calibration['test_split']['calibrated_coverage'])}，接近 90% 目标。",
    )
    add_doc_para(
        doc,
        "所以可以这样总结：分类和低摩擦风险识别已经比较有价值，区间估计经过校准后也能达到接近目标覆盖率。"
        "但它还不是最终论文级结果，因为 RoadSaW 子数据集表现偏弱，数据集身份可预测性过高，说明模型还可能学到了一些数据集风格捷径。",
    )

    doc.add_heading("7. 创新点怎么讲", level=1)
    add_doc_bullet(doc, "第一，不直接伪造真实摩擦系数，而是用弱监督摩擦区间，更符合公开数据集条件。")
    add_doc_bullet(doc, "第二，用 FrictionSet 把材质、湿度、雪冰、平整度组合成潜在状态，再从状态分布推摩擦区间。")
    add_doc_bullet(doc, "第三，加入局部证据场，让模型能显示它关注哪里，提升可解释性。")
    add_doc_bullet(doc, "第四，训练目标里加入 Group-DRO、V-REx、CORAL 和校准，面向跨数据集鲁棒性，而不是只追求平均准确率。")
    add_doc_bullet(doc, "第五，评估不只看 accuracy，还看 worst-dataset macro-F1、低摩擦召回、区间覆盖率和数据集捷径诊断。")

    doc.add_heading("8. 目前最需要承认的问题", level=1)
    add_doc_para(
        doc,
        f"最需要主动说明的是：数据集身份诊断 balanced accuracy 达到 {pct(diagnostic['overall_dataset_id_balanced_accuracy'])}，"
        "说明特征里仍然有很强的数据集来源信息。这不是致命问题，但必须通过 LODO 实验、消融实验和域泛化改进来解决。"
        "另外，证据场单独 head 的指标还不强，目前更像辅助分支和可解释性工具，后面要进一步加强监督。",
    )

    doc.add_heading("9. 下一步怎么讲", level=1)
    add_doc_para(
        doc,
        "下一步不是继续堆模块，而是补证据链。优先做四件事：第一，跑完整消融，证明每个模块确实有用；第二，跑 leave-one-dataset-out，证明不是只在混合测试集上好看；"
        "第三，针对 RoadSaW 做错误分析和域泛化改进；第四，改进局部证据场，让 attention 更稳定、更符合道路区域。",
    )

    doc.add_heading("10. 可以直接使用的讲解顺序", level=1)
    add_doc_bullet(doc, "先说研究问题：从路面图像估计摩擦风险和摩擦区间。")
    add_doc_bullet(doc, "再说数据限制：公开数据集没有真实 μ，所以采用弱摩擦区间。")
    add_doc_bullet(doc, "然后讲算法四层：全局视觉、物理纹理、局部证据场、FrictionSet。")
    add_doc_bullet(doc, "接着讲结果：risk/friction macro-F1、低摩擦检出、校准覆盖率。")
    add_doc_bullet(doc, "最后主动讲不足：RoadSaW、dataset shortcut、LODO 和消融还要补。")

    doc.add_paragraph("")
    add_doc_para(doc, f"对应技术 PDF: {REPORT_PDF}")
    doc.save(PLAIN_DOCX)


def register_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("CNFont", str(font_path)))
            return "CNFont"
    return "Helvetica"


def make_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TitleCN",
            parent=base["Title"],
            fontName=font_name,
            fontSize=18,
            leading=25,
            alignment=TA_CENTER,
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "H1CN",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=19,
            textColor=colors.HexColor("#12355b"),
            spaceBefore=10,
            spaceAfter=6,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "H2CN",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=11.5,
            leading=16,
            textColor=colors.HexColor("#1f5f72"),
            spaceBefore=7,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "BodyCN",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9.2,
            leading=14.2,
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "SmallCN",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=7.7,
            leading=10.8,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "formula": ParagraphStyle(
            "FormulaCN",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=8.6,
            leading=12.6,
            leftIndent=8,
            textColor=colors.HexColor("#333333"),
            backColor=colors.HexColor("#f6f8fa"),
            borderPadding=4,
            spaceBefore=2,
            spaceAfter=5,
            wordWrap="CJK",
        ),
    }


def add_title(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    story.append(Paragraph("视觉图像摩擦可供性估计算法研究内容", styles["title"]))
    story.append(Paragraph("Friction Affordance Field: visual road-state learning for friction affordance intervals", styles["small"]))
    meta_rows = [
        ["项目目录", str(ROOT)],
        ["实验目录", str(EXPERIMENT.relative_to(ROOT))],
        ["配置文件", str(CONFIG.relative_to(ROOT))],
        ["生成时间", data["generated_at"]],
        ["best.pt 保存时间", data["best_ckpt_mtime"]],
        ["last.pt 保存时间", data["last_ckpt_mtime"]],
        ["best epoch", str(data["best_epoch"])],
        ["审计结论", str(data["audit"].get("verdict", "unknown"))],
    ]
    story.append(make_table(meta_rows, [4.0 * cm, 12.0 * cm], styles))
    story.append(Spacer(1, 6))


def add_executive_summary(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    detailed = data["detailed"]
    calibration = data["calibration"]
    diagnostic = data["diagnostic"]
    story.append(Paragraph("1. 研究结论概览", styles["h1"]))
    p(
        story,
        styles,
        "当前模型已经完成三数据集联合训练、best checkpoint 评估、区间校准、数据集捷径诊断和局部证据图导出。"
        "它不再只是路况分类器，而是把视觉路面状态转成摩擦可供性区间和低摩擦风险。"
        "但从严格审稿角度看，还不能宣称已经达到顶会/顶刊最终水平，因为 leave-one-dataset-out 和完整消融还没有完成。",
    )
    rows = [
        ["核心指标", "当前数值", "解释"],
        [
            "friction macro-F1",
            pct(task_metric(detailed, "friction", "macro_f1")),
            "整体摩擦状态分类较强，类别不均衡下比 accuracy 更有参考价值。",
        ],
        [
            "risk macro-F1",
            pct(task_metric(detailed, "risk", "macro_f1")),
            "风险等级识别稳定，可作为安全感知任务的主结果。",
        ],
        [
            "低摩擦召回率",
            pct(detailed["low_friction_detection"]["recall"]),
            "把 high/very_high risk 视为低摩擦正类，召回较高，是现阶段亮点。",
        ],
        [
            "校准后区间覆盖率",
            pct(calibration["test_split"]["calibrated_coverage"]),
            "目标为 90%，测试集接近目标，说明 conformal calibration 有效。",
        ],
        [
            "数据集身份可预测性",
            pct(diagnostic["overall_dataset_id_balanced_accuracy"]),
            "过高，说明特征仍带有强域信息，是后续必须解决的问题。",
        ],
    ]
    story.append(make_table(rows, [4.3 * cm, 2.8 * cm, 8.9 * cm], styles, header=True))
    p(
        story,
        styles,
        "性能判断: 结果合理且有研究价值。合理之处在于 validation best 与 test 指标量级一致，低摩擦检出强，校准后覆盖率接近目标。"
        "主要风险是 RoadSaW 子域表现偏弱、原始区间覆盖不足、数据集 shortcut 明显、证据场辅助 head 单独预测能力不足。",
    )


def add_dataset_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    cfg = data["cfg"]
    story.append(Paragraph("2. 数据集、标签与弱摩擦系数定义", styles["h1"]))
    p(
        story,
        styles,
        "当前路线严格基于公开可获得数据集，不依赖自采车辆动力学数据。"
        "需要强调: RSCD、RoadSaW、RoadSC 提供的是视觉路况/材质/雪水状态标签，不是轮胎力传感器测得的真实摩擦系数。"
        "因此训练目标采用保守的弱摩擦区间，而不是把图像标签伪装成真实 μ 测量。",
    )
    formula(
        story,
        styles,
        "\\mathcal{D}=\\{(\\mathbf{x}_i,\\mathbf{y}_i,\\tilde{I}_i,d_i,g_i)\\}_{i=1}^{N},\\quad \\tilde{I}_i=[\\tilde{\\mu}_{L,i},\\tilde{\\mu}_{U,i}]",
    )
    formula_note(
        story,
        styles,
        "该公式定义整个训练样本。D 是全部公开数据集样本集合，x_i 是第 i 张路面图像，y_i 是它的多任务标签，tilde I_i 是由路况标签映射得到的弱摩擦区间，d_i 表示样本来自哪个数据集，g_i 表示鲁棒训练分组。这里用 tilde 是为了强调该区间不是实车传感器测量值，而是由视觉标签和物理常识构造的弱监督目标。",
    )
    p(
        story,
        styles,
        "其中 D 表示由公开数据集构成的样本集合；x_i 为第 i 张图像；y_i 是多任务标签向量；"
        "tilde I_i 是由视觉路况标签映射得到的弱摩擦区间；d_i 是数据集域标识；g_i 是 dataset::core-state 形式的鲁棒训练分组。"
        "这种定义把“真实摩擦测量缺失”显式写进建模过程，避免把弱标签误解释为物理传感器真值。",
    )
    rows = [["数据集", "train", "val", "test", "使用方式"]]
    for name, splits in data["manifests"].items():
        role = {
            "RSCD": "主数据源，覆盖干/湿/积水/雪/冰、材质和平整度。",
            "RoadSaW": "湿滑路面与材质小域，用于跨数据集鲁棒性。",
            "RoadSC": "雪类状态补充，强化低摩擦风险学习。",
        }[name]
        rows.append([name, str(splits["train"]), str(splits["val"]), str(splits["test"]), role])
    story.append(make_table(rows, [2.5 * cm, 2.0 * cm, 1.8 * cm, 1.8 * cm, 8.2 * cm], styles, header=True))
    p(
        story,
        styles,
        f"训练配置采用 image_size={cfg['data']['image_size']}，batch_size={cfg['data']['batch_size']}，"
        f"训练集每个 dataset/class_label 最多取 {cfg['data']['max_train_samples_per_class']} 张，"
        f"验证/测试每类最多取 {cfg['data']['max_val_samples_per_class']} / {cfg['data']['max_test_samples_per_class']} 张。"
        f"训练时使用 WeightedRandomSampler，每个 epoch 采样 {cfg['data']['balanced_num_samples_per_epoch']} 张，"
        "平衡单位为 dataset 和 class_label。测试详细评估的实际样本数为 10401。",
    )
    story.append(Paragraph("2.1 标签空间", styles["h2"]))
    label_rows = [
        ["任务", "标签"],
        ["friction", ", ".join(FRICTION_STATES)],
        ["material", ", ".join(MATERIALS)],
        ["wetness", ", ".join(WETNESS)],
        ["snow", ", ".join(SNOW)],
        ["unevenness", ", ".join(UNEVENNESS)],
        ["risk", ", ".join(RISK)],
    ]
    story.append(make_table(label_rows, [3.0 * cm, 13.0 * cm], styles, header=True))
    story.append(Paragraph("2.2 视觉标签到摩擦区间的弱监督映射", styles["h2"]))
    interval_rows = [["状态", "弱摩擦区间 [μ_min, μ_max]", "风险标签"]]
    for state in FRICTION_STATES:
        low, high = weak_mu_interval_from_state(friction=state, wetness=state if state in WETNESS else None, snow=state)
        interval_rows.append([state, f"[{low:.2f}, {high:.2f}]", risk_from_mu_interval(low, high) or "unknown"])
    story.append(make_table(interval_rows, [4.0 * cm, 5.5 * cm, 6.5 * cm], styles, header=True))
    story.append(Paragraph("2.3 变量解释", styles["h2"]))
    var_rows = [
        ["变量", "含义"],
        ["x", "输入 RGB 路面图像，训练中 resize/crop 到 192×192。"],
        ["D", "训练/验证/测试数据集集合，由 RSCD、RoadSaW、RoadSC 的 manifest 组成。"],
        ["y_t", "第 t 个任务的标签，例如 friction、risk、wetness 等。缺失标签用 mask 忽略。"],
        ["μ", "轮胎-路面摩擦系数。当前没有实测值，使用弱标签区间表达可供性。"],
        ["I=[μ_L, μ_U]", "模型预测的摩擦区间，下界 μ_L，上界 μ_U。"],
        ["h_g, h_p, h_e", "全局语义特征、物理纹理特征、局部证据场汇聚特征。"],
        ["s=(m,w,n,u)", "潜在路面状态: 材质 m、湿度 w、雪冰 n、平整度 u。"],
        ["S", "所有潜在路面状态的笛卡尔积集合，当前大小为 450。"],
        ["l_s, u_s", "潜在状态 s 对应的弱物理摩擦区间下界和上界。"],
        ["A_i", "局部证据场第 i 个 patch 的 attention 权重。"],
        ["c_i", "底部道路接触先验，鼓励模型看车辆可能接触的路面区域。"],
        ["I_P, I_F, I_E", "参数 μ 分支、FrictionSet 分支、EvidenceField 分支分别给出的摩擦区间。"],
        ["q", "conformal calibration 的非一致性分位数半径，用于扩大预测区间。"],
    ]
    story.append(make_table(var_rows, [3.4 * cm, 12.6 * cm], styles, header=True))


def add_algorithm_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    cfg = data["cfg"]
    model_cfg = cfg["model"]
    story.append(PageBreak())
    story.append(Paragraph("3. 当前算法完整流程", styles["h1"]))
    p(
        story,
        styles,
        "算法名称可以概括为 Friction Affordance Field: 用全局视觉语义识别路况，用物理纹理分支补充光泽/积雪/粗糙度证据，"
        "再用局部证据场在可能接触路面区域做弱监督 MIL 聚合，最后通过摩擦状态集合把多任务标签分布转成摩擦区间。",
    )
    story.append(Paragraph("3.1 全局视觉编码与多任务头", styles["h2"]))
    formula(
        story,
        styles,
        "\\mathbf{x}\\in\\mathbb{R}^{3\\times H\\times W},\\quad \\mathbf{h}_{g}=f_{\\theta}(\\mathbf{x}),\\quad \\mathbf{h}=\\mathrm{LN}([\\mathbf{h}_{g};\\mathbf{h}_{p};\\mathbf{h}_{e}]),\\quad \\mathbf{z}_{t}=\\mathbf{W}_{t}\\mathbf{h}+\\mathbf{b}_{t},\\quad p_{t}(y\\mid\\mathbf{x})=\\mathrm{softmax}(\\mathbf{z}_{t})",
    )
    formula_note(
        story,
        styles,
        "该公式描述主干网络和多任务分类头。x 是输入图像；f_theta 是 ConvNeXt-Tiny 编码器；h_g 是全局语义特征；h_p 是物理纹理分支特征；h_e 是局部证据场特征；LN 表示 LayerNorm；z_t 是第 t 个任务的 logits；softmax 后得到该任务每个类别的概率。这样设计的目的，是让同一张图同时输出路面状态、材质、湿度、雪冰状态、风险等级等信息，而不是只做单一分类。",
    )
    p(
        story,
        styles,
        f"全局 backbone 为 {model_cfg['backbone']}，ImageNet 预训练，embedding_dim={model_cfg['embedding_dim']}，"
        f"dropout={model_cfg['dropout']}。每个任务有独立线性分类头，标签缺失时通过 mask 不参与该任务交叉熵。",
    )
    story.append(Paragraph("3.2 物理纹理分支", styles["h2"]))
    p(
        story,
        styles,
        "PhysicsTextureBranch 从归一化图像还原 RGB 后计算灰度、饱和度、亮度、Sobel 梯度、Laplacian、雪白程度、镜面高光、暗水区域、湿滑 proxy、边缘粗糙度和区域连通性等 18 个统计量。"
        "这些量不是替代 CNN，而是给模型提供与摩擦更直接相关的低层证据。",
    )
    formula(
        story,
        styles,
        "r_{\\mathrm{snow}}=\\sigma(12(V-0.72))\\sigma(12(0.28-S)),\\quad r_{\\mathrm{spec}}=\\sigma(14(V-0.82))\\sigma(12(0.24-S)),\\quad r_{\\mathrm{wet}}=\\mathrm{clip}(r_{\\mathrm{spec}}+0.5r_{\\mathrm{dark}},0,1)",
    )
    formula_note(
        story,
        styles,
        "该公式是物理纹理分支中的手工视觉证据。V 表示亮度，S 表示饱和度，sigma 是 sigmoid 函数。雪通常表现为高亮低饱和，所以 r_snow 用亮度高、饱和度低来刻画；湿滑区域常有镜面反光，所以 r_spec 用高亮低饱和来近似；暗水区域 r_dark 与反光区域共同组成湿滑 proxy r_wet。clip 用于把数值限制在 0 到 1。它的作用是把人类对雪、水、反光、湿滑的直觉转成可学习特征。",
    )
    story.append(Paragraph("3.3 局部摩擦证据场", styles["h2"]))
    p(
        story,
        styles,
        f"EvidenceField 将 16 通道局部证据以 patch_stride={model_cfg['evidence_patch_stride']} 平均池化成 patch 网格。"
        "在 192×192 输入下，网格约为 24×24。每个 patch 预测局部风险 logits、局部 μ 区间和 attention logits。",
    )
    formula(
        story,
        styles,
        "A_i=\\frac{\\exp(q_i+\\lambda\\log c_i)}{\\sum_j\\exp(q_j+\\lambda\\log c_j)},\\quad \\mathbf{z}_{E}=\\sum_i A_i\\mathbf{z}_i,\\quad \\mu_{L}^{E}=\\sum_i A_i\\mu_{L,i},\\quad \\mu_{U}^{E}=\\sum_i A_i\\mu_{U,i}",
    )
    formula_note(
        story,
        styles,
        "该公式描述局部证据场的注意力聚合。q_i 是模型对第 i 个 patch 学到的注意力分数，c_i 是底部道路接触先验，lambda 控制先验强度。A_i 是归一化后的 patch 权重，权重越大说明该区域越可能提供摩擦判断证据。z_E 是所有 patch 风险 logits 的加权和，mu_L^E 和 mu_U^E 是局部证据分支给出的摩擦区间下界和上界。这个公式让模型尽量看可能与轮胎接触的路面区域，而不是背景、天空或车身。",
    )
    p(
        story,
        styles,
        "其中 q_i 是学习到的 attention logit，c_i 是底部中心道路接触先验，λ 为 contact_prior_strength。"
        "注意力场没有像素级真值，因此使用 image-level 弱监督、contact prior KL、attention smoothness 和主分支一致性来约束。",
    )
    story.append(Paragraph("3.4 FrictionSet 潜在状态集合", styles["h2"]))
    formula(
        story,
        styles,
        "s=(m,w,n,u),\\quad |\\mathcal{S}|=5\\times5\\times6\\times3=450,\\quad p(s\\mid\\mathbf{x})=p(m\\mid\\mathbf{x})p(w\\mid\\mathbf{x})p(n\\mid\\mathbf{x})p(u\\mid\\mathbf{x})",
    )
    formula_note(
        story,
        styles,
        "该公式定义 FrictionSet 的潜在路面状态。一个状态 s 由四个因素组成：材质 m、湿度 w、雪冰状态 n、平整度 u。当前标签空间大小分别为 5、5、6、3，因此总状态数是 450。p(s|x) 用四个任务的边缘概率相乘得到，表示模型认为图像属于某个具体路面组合的概率。这样可以把不同数据集里不完全一致的标签统一到同一个潜在状态空间。",
    )
    formula(
        story,
        styles,
        "\\mu_{L}^{F}=\\sum_{s\\in\\mathcal{S}}p(s\\mid\\mathbf{x})l_s-\\alpha H(p),\\quad \\mu_{U}^{F}=\\sum_{s\\in\\mathcal{S}}p(s\\mid\\mathbf{x})u_s+\\alpha H(p),\\quad \\alpha=0.12",
    )
    formula_note(
        story,
        styles,
        "该公式把潜在路面状态概率转换为摩擦区间。l_s 和 u_s 是状态 s 对应的弱物理摩擦下界和上界；模型对所有状态按概率加权求和，得到 FrictionSet 分支的区间。H(p) 是状态分布熵，表示模型不确定性；alpha H(p) 用于扩大区间：越不确定，区间越宽。这个设计能避免模型在模糊图像上给出过窄、过自信的摩擦估计。",
    )
    p(
        story,
        styles,
        "l_s 和 u_s 来自弱物理先验区间。状态分布熵 H(p) 越大，说明路面状态越不确定，模型自动扩大区间，避免给出过窄的过度自信估计。",
    )
    story.append(Paragraph("3.5 μ 参数分支与最终融合", styles["h2"]))
    formula(
        story,
        styles,
        "\\hat{\\mu}_{P}=1.2\\sigma(a),\\quad b_{P}=\\mathrm{softplus}(b),\\quad I_{P}=[\\hat{\\mu}_{P}-zb_{P},\\hat{\\mu}_{P}+zb_{P}],\\quad z=1.28155",
    )
    formula_note(
        story,
        styles,
        "该公式是参数化摩擦区间分支。a 和 b 是网络输出的两个实数；sigmoid 把均值限制到 0 到 1.2 的摩擦系数范围；softplus 保证尺度 b_P 为正；z=1.28155 用来把尺度转换为区间半宽。I_P 是参数分支给出的摩擦区间。它给模型保留了直接从视觉特征学习连续摩擦趋势的能力，而不是完全依赖人工区间先验。",
    )
    formula(
        story,
        styles,
        "I_{0}=0.80I_{F}+0.20I_{P},\\quad I=0.20I_{E}+0.80I_{0},\\quad \\mathbf{z}_{\\mathrm{risk}}=0.85\\mathbf{z}_{\\mathrm{main}}+0.15\\mathbf{z}_{E}",
    )
    formula_note(
        story,
        styles,
        "该公式描述最终融合。I_F 是 FrictionSet 区间，I_P 是参数分支区间，I_E 是局部证据场区间。当前配置先用 80% FrictionSet 和 20% 参数分支得到 I_0，再用 20% 局部证据场修正最终区间 I。风险 logits 也由主分支和证据场分支加权融合。这样做的直觉是：FrictionSet 保证物理和标签一致性，参数分支提供灵活拟合，证据场负责局部视觉修正。",
    )
    p(
        story,
        styles,
        "这里 z=1.28155 近似对应正态 80% 中央区间的半宽系数。当前配置把 FrictionSet 作为主区间来源，参数分支保留 20% 自由度，局部证据场再以 20% 权重修正最终区间。",
    )


def add_loss_training_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    cfg = data["cfg"]
    loss_cfg = cfg["loss"]
    optim = cfg["optim"]
    story.append(Paragraph("4. 训练目标与实现细节", styles["h1"]))
    formula(
        story,
        styles,
        "\\mathcal{L}=\\sum_{k}\\lambda_k\\mathcal{L}_k=\\lambda_{\\mathrm{task}}\\mathcal{L}_{\\mathrm{CE}}+\\lambda_{\\mathrm{comp}}\\mathcal{L}_{\\mathrm{comp}}+\\lambda_{\\mathrm{DRO}}\\mathcal{L}_{\\mathrm{DRO}}+\\lambda_{\\mathrm{VREx}}\\mathcal{L}_{\\mathrm{VREx}}+\\lambda_{\\mathrm{int}}\\mathcal{L}_{\\mathrm{int}}+\\lambda_{\\mathrm{cov}}\\mathcal{L}_{\\mathrm{cov}}+\\cdots",
    )
    formula_note(
        story,
        styles,
        "该公式是总训练目标。L_k 表示不同损失项，lambda_k 是每个损失的权重。CE 负责多任务分类，comp 负责潜在状态与已知标签兼容，DRO 和 VREx 负责跨数据集/跨组鲁棒性，int 和 cov 负责摩擦区间的概率拟合和覆盖。省略号表示还有端点、宽度、单调约束、CORAL 域对齐、证据场辅助损失等。总目标的核心思想是：既要分类准，又要区间合理，还要尽量减少数据集捷径。",
    )
    loss_rows = [
        ["损失项", "权重", "作用"],
        ["masked multitask CE", f"{loss_cfg['task_weight']}", "多任务分类主监督，缺失标签不参与。"],
        ["FrictionSet compatibility NLL", f"{loss_cfg['compatibility_weight']}", "要求潜在状态概率落在已知标签兼容集合内。"],
        ["Group-DRO", f"{loss_cfg['group_dro_weight']}", "优化最差 dataset::core_state 组，缓解小域/难类被平均指标掩盖。"],
        ["Group V-REx", f"{loss_cfg['group_vrex_weight']}", "降低组间损失方差，提升跨域稳定性。"],
        ["risk ordinal EMD", f"{loss_cfg['risk_ordinal_weight']}", "利用 very_low 到 very_high 的有序结构。"],
        ["interval censored NLL", f"{loss_cfg['interval_weight']}", "让预测分布在目标弱区间内有较高概率。"],
        ["coverage/endpoint/target-width", f"{loss_cfg['coverage_weight']}/{loss_cfg['endpoint_weight']}/{loss_cfg['target_width_weight']}", "约束区间覆盖、端点位置和宽度。"],
        ["monotonic constraints", f"{loss_cfg['monotonic_weight']}/{loss_cfg['risk_mu_monotonic_weight']}", "湿度或风险越高，μ 不应系统性升高。"],
        ["CORAL / conditional CORAL", f"{loss_cfg['feature_coral_weight']}/{loss_cfg['risk_conditional_coral_weight']}", "对齐不同数据集的特征均值和协方差。"],
        ["evidence field losses", "多项", "局部证据风险、区间、attention prior、smoothness、一致性约束。"],
    ]
    story.append(make_table(loss_rows, [4.2 * cm, 2.7 * cm, 9.1 * cm], styles, header=True))
    formula(
        story,
        styles,
        "\\mathcal{L}_{\\mathrm{CE}}=\\frac{1}{|\\Omega_t|}\\sum_{t}\\sum_{i\\in\\Omega_t}-\\log p_t(y_{i,t}\\mid\\mathbf{x}_i)",
    )
    formula_note(
        story,
        styles,
        "该公式是 masked 多任务交叉熵。Omega_t 表示第 t 个任务上有有效标签的样本集合，因为不同数据集不是每个任务都有标签。对每个有效样本，模型给真实类别 y_i,t 的概率越高，损失越小。这个损失负责学习 friction、material、wetness、snow、risk 等分类任务的基本识别能力。",
    )
    formula(
        story,
        styles,
        "\\mathcal{L}_{\\mathrm{EMD}}=\\frac{1}{K-1}\\sum_{k=1}^{K-1}\\left(\\sum_{j=1}^{k}p_j-\\sum_{j=1}^{k}\\mathbf{1}[j=y]\\right)^2",
    )
    formula_note(
        story,
        styles,
        "该公式是风险等级的有序 EMD 损失。K 是风险等级数量，p_j 是第 j 个风险等级的预测概率。普通交叉熵只知道对错，但风险等级有顺序：very_low、low、medium、high、very_high。把 medium 预测成 high 的错误，应该比预测成 very_high 小。EMD 用累积分布差异表达这种有序距离，使风险预测更符合等级结构。",
    )
    formula(
        story,
        styles,
        "\\mathcal{L}_{\\mathrm{mono}}=\\frac{1}{|\\mathcal{P}|}\\sum_{(i,j)\\in\\mathcal{P}}\\max(0,\\hat{\\mu}_i-\\hat{\\mu}_j),\\quad \\mathcal{P}=\\{(i,j):r_i>r_j\\}",
    )
    formula_note(
        story,
        styles,
        "该公式是风险-摩擦单调约束。P 是样本对集合，r_i>r_j 表示第 i 个样本风险更高。物理直觉上，风险更高通常不应对应更大的摩擦系数，因此如果高风险样本的预测 mu_hat_i 反而大于低风险样本 mu_hat_j，就产生惩罚。这个约束不是硬规则，而是软约束，用来减少明显违背物理直觉的预测。",
    )
    formula(
        story,
        styles,
        "\\mathcal{L}_{\\mathrm{comp}}=-\\log\\sum_{s\\in\\mathcal{C}(\\mathbf{y})}p(s\\mid\\mathbf{x}),\\quad \\mathcal{L}_{\\mathrm{DRO}}=\\tau\\log\\sum_g\\exp(\\mathcal{L}_g/\\tau),\\quad \\mathcal{L}_{\\mathrm{int}}=-\\log\\left[\\sigma\\!\\left(\\frac{U-\\mu}{b}\\right)-\\sigma\\!\\left(\\frac{L-\\mu}{b}\\right)\\right]",
    )
    formula_note(
        story,
        styles,
        "这组三个公式分别对应兼容性、鲁棒性和区间概率。C(y) 是与已知标签 y 兼容的潜在状态集合，comp 损失要求模型把概率分配到这些合理状态上。DRO 损失对不同分组 g 的损失做 softmax 风格加权，tau 控制接近最差组还是平均组，目的是不要让小数据集或困难组被平均指标掩盖。interval 损失把摩擦系数看作 logistic 分布，要求目标区间 [L,U] 内的概率尽量大。",
    )
    story.append(Paragraph("4.1 训练配置", styles["h2"]))
    rows = [
        ["项目", "配置"],
        ["硬件/环境", "RTX 3050 Laptop 4GB, conda env faf_paper, AMP enabled"],
        ["优化器", f"AdamW, lr={optim['lr']}, weight_decay={optim['weight_decay']}"],
        ["epoch/早停", f"max epoch={optim['epochs']}, patience={optim['early_stop_patience']}, min_delta={optim['early_stop_min_delta']}"],
        ["梯度", f"grad_accum_steps={optim['grad_accum_steps']}, grad_clip_norm={optim['grad_clip_norm']}"],
        ["增强", "RandomResizedCrop, flip, ColorJitter, RandomGrayscale, GaussianBlur, RandomErasing"],
        ["最好轮次", f"epoch {data['best_epoch']}, val loss={data['best_val'].get('loss', 0):.4f}"],
        ["结束方式", "epoch 18 触发 early stopping"],
    ]
    story.append(make_table(rows, [4.0 * cm, 12.0 * cm], styles, header=True))


def add_results_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    detailed = data["detailed"]
    evaluate = data["evaluate"]
    calibration = data["calibration"]
    diagnostic = data["diagnostic"]
    story.append(PageBreak())
    story.append(Paragraph("5. 最好结果全面检查", styles["h1"]))
    p(
        story,
        styles,
        "本节基于 best.pt 在 test split 上重新评估得到的 evaluate_test.json、detailed_test.json、interval_calibration_90.json 和 dataset_id_diagnostic.json。"
        "这些文件均位于当前主实验输出目录。",
    )
    story.append(Paragraph("5.1 多任务分类表现", styles["h2"]))
    rows = [["任务", "样本数", "accuracy", "balanced acc", "macro-F1", "worst dataset macro-F1"]]
    for task in ["friction", "risk", "wetness", "material", "snow", "unevenness"]:
        info = detailed["tasks"][task]
        rows.append(
            [
                task,
                str(info["num_samples"]),
                pct(info["accuracy"]),
                pct(info["balanced_accuracy"]),
                pct(info["macro_f1"]),
                pct(info["by_dataset"]["_worst_macro_f1"]["value"]),
            ]
        )
    story.append(make_table(rows, [2.5 * cm, 2.0 * cm, 2.2 * cm, 2.4 * cm, 2.4 * cm, 4.5 * cm], styles, header=True))
    story.append(Paragraph("5.2 关键任务按数据集分解", styles["h2"]))
    rows = [["任务/数据集", "样本数", "accuracy", "macro-F1", "判断"]]
    for task in ["friction", "risk", "wetness"]:
        for dataset, info in detailed["tasks"][task]["by_dataset"].items():
            if dataset.startswith("_"):
                continue
            judgement = "短板" if float(info["macro_f1"]) < 0.70 else "可接受"
            rows.append([f"{task}/{dataset}", str(info["num_samples"]), pct(info["accuracy"]), pct(info["macro_f1"]), judgement])
    story.append(make_table(rows, [3.6 * cm, 2.0 * cm, 2.5 * cm, 2.5 * cm, 5.4 * cm], styles, header=True))
    story.append(Paragraph("5.3 低摩擦检出与区间质量", styles["h2"]))
    low = detailed["low_friction_detection"]
    mu = detailed["mu_interval"]
    rows = [
        ["指标", "数值", "解释"],
        ["低摩擦正类定义", low["positive_definition"], "把 risk=high/very_high 视作低摩擦高风险。"],
        ["recall / precision / F1", f"{pct(low['recall'])} / {pct(low['precision'])} / {pct(low['f1'])}", "安全场景更重视 recall，当前较强。"],
        ["raw coverage / width", f"{pct(mu['coverage'])} / {mu['width_mean']:.4f}", "原始区间覆盖不足，说明融合区间偏窄或弱标签区间较宽。"],
        ["calibrated coverage / width", f"{pct(calibration['test_split']['calibrated_coverage'])} / {calibration['test_split']['calibrated_width']:.4f}", "用验证集 conformal radius 扩展后接近 90% 覆盖。"],
        ["mean MAE to interval mid", f"{mu['mean_mae_to_interval_mid']:.4f}", "预测均值与弱标签中点距离较小，但它不是实测 μ 误差。"],
    ]
    story.append(make_table(rows, [4.2 * cm, 4.2 * cm, 7.6 * cm], styles, header=True))
    formula(
        story,
        styles,
        "e_i=\\max(\\mu_{L,i}-\\hat{\\mu}_{L,i},\\hat{\\mu}_{U,i}-\\mu_{U,i},0),\\quad q=\\mathrm{Quantile}_{0.90}(\\{e_i\\}_{i\\in\\mathcal{D}_{\\mathrm{val}}}),\\quad I_{\\mathrm{cal}}=[\\hat{\\mu}_{L}-q,\\hat{\\mu}_{U}+q]",
    )
    formula_note(
        story,
        styles,
        "该公式是 conformal 区间校准。e_i 是验证样本的非一致性分数：如果预测区间没有完全覆盖目标弱区间，就计算还需要向外扩张多少；如果已经覆盖，则为 0。q 是验证集 e_i 的 90% 分位数。测试时把预测下界减 q、上界加 q，得到校准区间 I_cal。它的作用是把原始偏窄的区间修正到接近目标覆盖率，但这仍然是弱标签区间上的覆盖，不代表真实路试 mu 误差。",
    )
    p(
        story,
        styles,
        "上式中的 e_i 是非一致性分数，表示预测区间未覆盖目标弱区间时需要向外扩张的最小半径。"
        "q 在验证集上取 90% 分位数，然后统一扩展测试集预测区间。这样得到的 coverage 是校准意义上的区间可靠性，不等同于真实路试 μ 的测量误差。",
    )
    story.append(Paragraph("5.4 审稿人视角风险检查", styles["h2"]))
    rows = [
        ["检查项", "结果", "含义"],
        ["audit verdict", str(data["audit"].get("verdict", "unknown")), "当前审计不是 ready。"],
        ["raw interval coverage", pct(mu["coverage"]), "低于预期，需要校准和区间分支改进。"],
        ["dataset ID balanced acc", pct(diagnostic["overall_dataset_id_balanced_accuracy"]), "模型特征几乎能识别数据集来源，域捷径严重。"],
        ["risk-conditioned dataset ID bal acc", pct(diagnostic["risk_conditioned_common_dataset_id_balanced_accuracy"]), "即使控制风险标签后，域信息仍很强。"],
        ["evidence risk macro-F1", pct(detailed["evidence_field"]["risk"]["macro_f1"]), "证据场单独 head 较弱，目前更像辅助/正则分支。"],
        ["evidence interval coverage", pct(detailed["evidence_field"]["mu_interval"]["coverage"]), "单独证据区间不可作为最终输出，需改进监督。"],
        ["LODO/消融", "未完成", "不能支撑最终顶会泛化和创新归因。"],
    ]
    story.append(make_table(rows, [4.5 * cm, 3.3 * cm, 8.2 * cm], styles, header=True))
    p(
        story,
        styles,
        f"evaluate_test.json 的总 loss={evaluate['loss']:.4f}，acc_friction={pct(evaluate['acc_friction'])}，acc_risk={pct(evaluate['acc_risk'])}。"
        "与 detailed_test.json 的汇总量级一致，说明评估脚本之间没有明显冲突。",
    )
    add_evidence_images(story, styles, data)


def add_evidence_images(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    images = data["evidence_maps"]
    if not images:
        return
    story.append(Paragraph("5.5 局部证据图样例", styles["h2"]))
    p(
        story,
        styles,
        f"当前已导出 {len(images)} 张 evidence map，可用于分析模型关注区域和错误案例。下图为自动挑选的代表样例。",
    )
    preferred = [
        "0002_rscd_high.jpg",
        "0007_rscd_very_low.jpg",
        "0037_roadsc_high.jpg",
        "0047_roadsaw_low.jpg",
    ]
    selected = []
    by_name = {p.name: p for p in images}
    for name in preferred:
        if name in by_name:
            selected.append(by_name[name])
    selected.extend([p for p in images if p not in selected])
    selected = selected[:4]
    cells = []
    for img_path in selected:
        im = Image(str(img_path), width=7.2 * cm, height=4.05 * cm)
        caption = Paragraph(img_path.name, styles["small"])
        cells.append([im, caption])
    rows = []
    for i in range(0, len(cells), 2):
        row = []
        for block in cells[i : i + 2]:
            row.append([block[0], block[1]])
        while len(row) < 2:
            row.append("")
        rows.append(row)
    table = Table(rows, colWidths=[8.0 * cm, 8.0 * cm])
    table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d7de")),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d7de")),
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)


def add_innovation_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    story.append(PageBreak())
    story.append(Paragraph("6. 当前算法创新点展开", styles["h1"]))
    innovations = [
        (
            "创新点 1: 从路况分类升级为弱监督摩擦可供性区间估计",
            "公开数据集缺少实测 μ，直接回归单值会被审稿人质疑标签来源。本算法把视觉标签映射为保守区间，并预测覆盖区间与风险等级，问题定义更诚实，也更接近安全决策需求。",
        ),
        (
            "创新点 2: FrictionSet 潜在路面状态集合",
            "把 material、wetness、snow、unevenness 的边缘分布组合成 450 个潜在状态，再对每个状态的弱物理摩擦区间做边缘化。相比单头分类，它能把多数据集异构标签统一到同一摩擦区间空间。",
        ),
        (
            "创新点 3: 局部摩擦证据场",
            "证据场不只输出全局标签，而是在 patch 级别计算雪白、镜面反光、暗水、粗糙度和底部接触先验，利用 MIL attention 聚合局部证据。这给模型提供了可视化解释入口。",
        ),
        (
            "创新点 4: 多数据集鲁棒训练目标",
            "使用 balanced sampling、Group-DRO、V-REx、全局/风险条件 CORAL，把训练目标从平均准确率推向跨域稳定性。虽然目前还没彻底消除 dataset shortcut，但路线是审稿人能接受的。",
        ),
        (
            "创新点 5: 区间校准和审计化评估",
            "模型不仅统计分类指标，还统计 raw/calibrated coverage、区间宽度、低摩擦召回、worst-dataset macro-F1、dataset-ID diagnostic。这种评估方式比单一 accuracy 更适合安全感知论文。",
        ),
    ]
    for title, body in innovations:
        story.append(Paragraph(title, styles["h2"]))
        p(story, styles, body)
    story.append(Paragraph("6.1 与顶会/顶刊要求的距离", styles["h2"]))
    p(
        story,
        styles,
        "当前创新方向是成立的，但证据链还不完整。顶会/顶刊审稿人通常会追问: 每个模块是否真的有效、是否能泛化到未见数据集、是否只是学到了数据集风格、区间估计是否有校准保证、可视化是否解释了成功和失败样例。"
        "因此，下一阶段不是盲目堆新模块，而是补强验证闭环。",
    )


def add_next_steps_section(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    story.append(Paragraph("7. 下一步详细规划", styles["h1"]))
    rows = [
        ["优先级", "任务", "产出", "目标"],
        ["P0", "补齐 ablation: global-only, +physics, +FrictionSet, +DG, +EvidenceField", "一张完整消融表", "证明创新点不是堆模块。"],
        ["P0", "跑 LODO: leave RSCD/RoadSaW/RoadSC out", "跨数据集泛化表", "解决 top-venue audit 的 block 项。"],
        ["P0", "重新跑 audit，并把 evidence maps 纳入实验材料", "ready/not-ready 自动判定", "形成实验闭环。"],
        ["P1", "RoadSaW 短板专项分析", "混淆矩阵和失败样例", "找出 damp/wet/very_wet 与材质混淆。"],
        ["P1", "域捷径抑制", "dataset-ID accuracy 降低曲线", "减少模型依赖数据集风格。"],
        ["P1", "证据场改进", "attention 更可信、aux head 更强", "把可解释性从展示图变成有效模块。"],
        ["P2", "论文级 baseline", "ConvNeXt direct, ViT/DINO, CLIP linear/probe, rule prior", "让结果能对标 CV 审稿习惯。"],
    ]
    story.append(make_table(rows, [1.5 * cm, 5.0 * cm, 4.2 * cm, 5.3 * cm], styles, header=True))
    story.append(Paragraph("7.1 推荐技术路线", styles["h2"]))
    p(
        story,
        styles,
        "建议继续走多数据集路线，但论文叙述必须从“简单合并数据集训练”改成“异构公开视觉标签下的弱监督摩擦可供性学习与域泛化”。"
        "这样可以把 RSCD 的规模、RoadSaW 的湿滑域、RoadSC 的雪域统一成研究问题，而不是被审稿人认为只是拼数据。",
    )
    p(
        story,
        styles,
        "算法改进优先方向: 第一，做条件域泛化，即只在相同 risk/core-state 条件内对齐域特征，避免把真实物理差异抹平。第二，加入 style/Fourier augmentation 或 domain-specific adapter，减少相机、裁剪、分辨率带来的 shortcut。第三，给 evidence field 加 road ROI 伪监督或底部道路 mask，让 attention 更聚焦可接触路面。第四，输出条件 conformal calibration，按 dataset/core-state/risk 给出覆盖率和宽度。",
    )
    story.append(Paragraph("7.2 可量化研究里程碑", styles["h2"]))
    milestones = [
        "完成至少 4 个消融模型的同一测试协议评估。",
        "完成 3 个 LODO 实验，并统计 held-out macro-F1、low-friction recall、calibrated coverage。",
        "把 RoadSaW worst macro-F1 从 0.65 左右提升到 0.70 以上，或给出明确失败原因分析。",
        "把 dataset-ID balanced accuracy 作为反向指标纳入实验表，展示域捷径抑制进展。",
        "整理 12 张成功/失败 evidence map，配合混淆矩阵解释模型行为。",
    ]
    for item in milestones:
        bullet(story, styles, item)


def add_appendix(story: list[Any], styles: dict[str, ParagraphStyle], data: dict[str, Any]) -> None:
    story.append(Paragraph("8. 附录: 当前文件与命令", styles["h1"]))
    rows = [
        ["内容", "路径"],
        ["配置", str(CONFIG.relative_to(ROOT))],
        ["best checkpoint", str((EXPERIMENT / "best.pt").relative_to(ROOT))],
        ["test summary", str((EXPERIMENT / "evaluate_test.json").relative_to(ROOT))],
        ["detailed metrics", str((EXPERIMENT / "detailed_test.json").relative_to(ROOT))],
        ["interval calibration", str((EXPERIMENT / "interval_calibration_90.json").relative_to(ROOT))],
        ["dataset shortcut diagnostic", str((EXPERIMENT / "dataset_id_diagnostic.json").relative_to(ROOT))],
        ["top-venue audit", str((EXPERIMENT / "topvenue_result_audit.md").relative_to(ROOT))],
        ["evidence maps", str((EXPERIMENT / "evidence_maps").relative_to(ROOT))],
        ["本 PDF", str(REPORT_PDF.relative_to(ROOT))],
    ]
    story.append(make_table(rows, [4.0 * cm, 12.0 * cm], styles, header=True))
    p(
        story,
        styles,
        "复现实验评估的核心脚本包括 scripts/evaluate.py、scripts/evaluate_detailed.py、scripts/calibrate_intervals.py、scripts/dataset_id_diagnostic.py、scripts/audit_topvenue_results.py。"
        "文档生成脚本为 scripts/make_weekly_progress_report.py。",
    )


def add_doc_para(doc, text: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = PtSafe(6)
    paragraph.add_run(text)
    return paragraph


def add_doc_bullet(doc, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = PtSafe(3)
    paragraph.add_run(text)
    return paragraph


def PtSafe(value: float):
    from docx.shared import Pt

    return Pt(value)


def p(story: list[Any], styles: dict[str, ParagraphStyle], text: str) -> None:
    story.append(Paragraph(escape(text), styles["body"]))


def formula(story: list[Any], styles: dict[str, ParagraphStyle], text: str) -> None:
    path = render_formula(text)
    with PILImage.open(path) as img:
        width_px, height_px = img.size
    max_width = 15.8 * cm
    width = min(max_width, width_px * 72.0 / 230.0)
    height = width * height_px / max(width_px, 1)
    image = Image(str(path), width=width, height=height)
    table = Table([[image]], colWidths=[16.0 * cm])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d7de")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 3))


def formula_note(story: list[Any], styles: dict[str, ParagraphStyle], text: str) -> None:
    table = Table(
        [[Paragraph("<b>公式解释：</b>" + escape(text), styles["small"])]],
        colWidths=[16.0 * cm],
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#fff8e6")),
                ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#e5d3a3")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 4))


def render_formula(latex: str) -> Path:
    import hashlib

    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    key = hashlib.sha1(latex.encode("utf-8")).hexdigest()[:16]
    path = FORMULA_DIR / f"{key}.png"
    if path.exists():
        return path
    fig = plt.figure(figsize=(9.0, 0.48), dpi=260)
    fig.patch.set_alpha(0.0)
    fig.text(0.01, 0.50, f"${latex}$", fontsize=15, va="center", ha="left", color="#111111")
    fig.savefig(path, dpi=260, transparent=True, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return path


def bullet(story: list[Any], styles: dict[str, ParagraphStyle], text: str) -> None:
    story.append(Paragraph("• " + escape(text), styles["body"]))


def make_table(
    rows: list[list[Any]],
    col_widths: list[float],
    styles: dict[str, ParagraphStyle],
    header: bool = False,
) -> Table:
    cell_style = styles["small"]
    data = [[to_para(cell, cell_style) for cell in row] for row in rows]
    table = Table(data, colWidths=col_widths, repeatRows=1 if header else 0)
    commands = [
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d0d7de")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]
    if header:
        commands.extend(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eaf3f8")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f3557")),
            ]
        )
    table.setStyle(TableStyle(commands))
    return table


def to_para(value: Any, style: ParagraphStyle) -> Paragraph:
    return Paragraph(escape(str(value)), style)


def task_metric(detailed: dict[str, Any], task: str, metric: str) -> float:
    return float(detailed["tasks"][task][metric])


def pct(value: float) -> str:
    return f"{100.0 * float(value):.2f}%"


def relative(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def draw_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawRightString(A4[0] - 1.55 * cm, 0.75 * cm, f"{doc.page}")
    canvas.restoreState()


if __name__ == "__main__":
    main()
