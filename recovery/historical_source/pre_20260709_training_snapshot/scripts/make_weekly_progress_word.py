from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
EXP_ROOT = Path(r"D:\NMI_SPWFM_datasets\friction_affordance_outputs\paper_protocol")
SUMMARY = PROJECT / "reports" / "paper_protocol_summary"
DESKTOP = Path(r"C:\Users\DELL\Desktop")


def main() -> None:
    report = build_report()
    out = DESKTOP / "视觉道路摩擦可供性估计研究进展.docx"
    doc = Document()
    setup_document(doc)
    write_document(doc, report)
    try:
        doc.save(out)
    except PermissionError:
        out = DESKTOP / "视觉道路摩擦可供性估计研究进展_更新版.docx"
        doc.save(out)
    print(out)


def build_report() -> dict[str, Any]:
    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "queue": read_json(SUMMARY / "queue_recovery_report.json"),
        "closure": read_json(SUMMARY / "ten_hour_closure_report.json"),
        "roadsc_state": read_json(EXP_ROOT / "baseline_single_roadsc_global_convnext" / "training_state.json"),
        "roadsc_detail": read_json(EXP_ROOT / "baseline_single_roadsc_global_convnext" / "detailed_test.json"),
        "roadsc_eval": read_json(EXP_ROOT / "baseline_single_roadsc_global_convnext" / "evaluate_test.json"),
        "fair_rows": collect_fair_rows(),
        "dataset_rows": [
            ["RSCD", "958,941 / 19,860 / 49,500", "路面状态、湿度、风险等弱标签", "样本最多，是 RSCD 单数据集公平对比和主要鲁棒性检查对象。"],
            ["RoadSaW", "15,360 / 4,812 / 1,728", "道路表面、湿滑/天气相关标签", "验证湿滑、材质和风险可供性；高反光/近白样本是难点。"],
            ["RoadSC", "4,374 / 1,467 / 573", "雪、冰、湿滑相关状态", "规模较小，适合做雪冰低摩擦场景补充验证。"],
        ],
    }


def collect_fair_rows() -> list[dict[str, Any]]:
    pairs = [
        ("RoadSaW", "single_roadsaw_full_faf", "baseline_single_roadsaw_global_convnext"),
        ("RSCD", "single_rscd_full_faf", "baseline_single_rscd_global_convnext"),
        ("RoadSC", "single_roadsc_full_faf", "baseline_single_roadsc_global_convnext"),
    ]
    rows = []
    for dataset, faf_name, base_name in pairs:
        faf = read_json(EXP_ROOT / faf_name / "detailed_test.json")
        base = read_json(EXP_ROOT / base_name / "detailed_test.json")
        rows.append(
            {
                "dataset": dataset,
                "faf_status": status_from(faf, EXP_ROOT / faf_name),
                "base_status": status_from(base, EXP_ROOT / base_name),
                "faf_friction": task_metric(faf, "friction", "macro_f1"),
                "base_friction": task_metric(base, "friction", "macro_f1"),
                "faf_risk": task_metric(faf, "risk", "macro_f1"),
                "base_risk": task_metric(base, "risk", "macro_f1"),
                "faf_low": low_recall(faf),
                "base_low": low_recall(base),
                "faf_cov": mu_metric(faf, "coverage"),
                "base_cov": mu_metric(base, "coverage"),
                "faf_width": mu_metric(faf, "width_mean"),
                "base_width": mu_metric(base, "width_mean"),
            }
        )
    return rows


def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [("Title", 18, "1F4E79"), ("Heading 1", 14, "1F4E79"), ("Heading 2", 12, "365F91")]:
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def write_document(doc: Document, report: dict[str, Any]) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("视觉道路摩擦可供性估计研究进展")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    p = doc.add_paragraph(f"生成时间：{report['generated_at']}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_h1(doc, "零、汇报摘要")
    add_para(
        doc,
        "一句话概括：本研究不是直接回归实测轮胎-路面摩擦系数，而是基于公开道路图像标签进行弱监督视觉摩擦可供性区间估计，"
        "核心价值在于把道路状态识别推进到风险、区间和局部证据解释。",
    )
    add_numbered(
        doc,
        [
            "已经完成 P0 核心闭环：消融实验、LODO 跨数据集压力测试、RoadSaW/RSCD 同数据集强基线公平对比、RSCD 基线 bootstrap 和审计。",
            "当前结果表明 PhysicsTexture 是有效模块，RoadSaW 上 FAF 相对强 ConvNeXt 基线有小幅收益；但 RSCD 上强基线超过完整 FAF，说明复杂模块存在负迁移，必须裁剪。",
            "下一步最合理路线不是继续盲目堆模块，而是转向 segmentation-style 局部道路证据、masked consistency、区间校准和域偏移诊断，把模型学到的数据集捷径压下去。",
        ],
    )
    add_para(
        doc,
        "汇报时建议避免的表述：不要说“已经实现真实摩擦系数估计”或“多数据集泛化已经成功”。更准确的说法是："
        "当前完成了公开图像标签下的弱监督摩擦可供性估计，发现了跨数据集泛化失败和强基线挑战，并据此确定了下一代算法改进方向。",
    )

    add_h1(doc, "一、研究对象和边界")
    add_para(
        doc,
        "本研究面向基于公开道路图像数据集的视觉摩擦可供性估计。需要严格区分："
        "现有 RSCD、RoadSaW、RoadSC 等公开数据集没有同步实测轮胎-路面摩擦系数，"
        "因此当前研究不能表述为真实摩擦系数直接回归，而应表述为基于公开视觉标签的弱监督摩擦等级和摩擦区间估计。",
    )
    add_para(
        doc,
        "核心问题是：给定一张道路图像，模型不仅要判断路面是否干燥、湿滑、积雪或结冰，"
        "还要把这些视觉状态映射为摩擦风险等级和摩擦可供性区间，并尽量让判断来自路面区域本身，"
        "而不是数据集风格、相机裁剪、背景颜色等捷径。",
    )

    add_h1(doc, "二、数据集和标签使用")
    add_table(doc, ["数据集", "训练/验证/测试", "使用标签", "作用"], report["dataset_rows"])
    add_para(
        doc,
        "标签映射原则：干燥沥青或混凝土通常映射到高摩擦区间；湿路、积水、雪、冰、泥泞映射到较低摩擦区间；"
        "材质、粗糙度、湿度、雪冰标签作为辅助任务，用来解释低摩擦来源。这样做的优点是完全基于可获得公开数据，"
        "缺点是只能得到弱监督区间，而不是传感器实测摩擦系数。",
    )
    add_h1(doc, "二点五、外部资料对数据差异的解释")
    add_para(
        doc,
        "RSCD 官方说明其 friction、material、unevenness 三类属性用于底盘控制和驾驶辅助，friction 包含 dry、wet、water、fresh snow、melted snow、ice 六类；"
        "数据发布页还说明原始图像来自车载相机，但发布图像是只包含道路表面区域的 240×360 patch。因此 RSCD 更像局部道路表面裁剪，而不是完整前视驾驶图。",
    )
    add_para(
        doc,
        "RoadSaW 官方说明其数据是由标定相机生成的鸟瞰路面 patches，包含 asphalt、cobblestone、concrete 三种表面和四级湿度；"
        "湿度标注来自准确测量。由此可解释 RoadSaW 中许多近白图像：它们很可能对应湿路、反光、水膜、光照和鸟瞰变换后的高亮区域，"
        "不能简单当作坏图删除，而应作为湿滑/反光难例处理。",
    )
    add_para(
        doc,
        "RoadSC 与 RoadSaW 兼容，主要补充雪覆盖表面。由于 RoadSC/RoadSaW 的视角、patch 生成方式、分辨率和采集环境与 RSCD 不同，"
        "它们与 RSCD 的图像风格差异是合理现象，也是 LODO 泛化失败的重要原因。",
    )
    add_para(
        doc,
        "近期文献也支持这个判断：视觉道路摩擦估计通常分为两类，一类是用公开图像标签做 road condition / wetness / surface proxy，"
        "另一类是用专门采集的光学摩擦传感器或车辆动力学真值做连续摩擦回归。我们的公开数据路线应严谨定位为前者，除非引入带真实摩擦传感器的数据。",
    )

    add_h1(doc, "三、算法整体流程")
    add_para(doc, "算法可以理解为“全局语义 + 视觉物理纹理 + 多任务弱监督 + 区间校准”的流程。")
    add_numbered(
        doc,
        [
            "输入道路图像 x，经过统一尺寸、归一化和数据增强后送入 ConvNeXt 主干网络。",
            "主干网络提取全局语义特征 h = fθ(x)，用于判断整体道路状态。",
            "PhysicsTexture 分支提取湿滑、反光、低纹理、粗糙度等视觉物理线索 p = gφ(x)。",
            "融合特征 z = concat(h, p)，再输入多个任务头，同时预测摩擦等级、风险等级、湿度、雪冰、材质、平整度等标签。",
            "区间头输出 Iμ(x) = [μmin(x), μmax(x)]，表示该图像对应的弱摩擦可供性范围。",
            "训练后再用验证集进行 conformal calibration，把原始区间扩张到目标覆盖率附近，从而报告校准覆盖率和区间宽度。",
        ],
    )

    add_h1(doc, "四、公式和变量解释")
    add_formula(doc, "h = f_θ(x)")
    add_para(doc, "其中 x 表示输入道路图像，fθ 表示 ConvNeXt 主干网络，θ 是主干网络参数，h 是全局视觉语义特征。")
    add_formula(doc, "p = g_φ(x)")
    add_para(doc, "其中 gφ 表示 PhysicsTexture 分支，φ 是该分支参数，p 表示显式视觉物理纹理特征，例如反光、粗糙、低纹理和湿滑线索。")
    add_formula(doc, "z = concat(h, p)")
    add_para(doc, "其中 z 是融合后的特征，concat 表示特征拼接。若某个基线不使用 PhysicsTexture，则 z 退化为全局特征 h。")
    add_formula(doc, "I_μ(x) = [ μ_min(x), μ_max(x) ]")
    add_para(doc, "其中 Iμ(x) 是模型输出的弱摩擦区间，μmin(x) 和 μmax(x) 分别表示该图像对应摩擦可供性的下界和上界。")
    add_formula(doc, "L = L_cls + λ_aux L_aux + λ_int L_interval + λ_mono L_mono + λ_cal L_cal")
    add_para(
        doc,
        "总损失 L 由五部分组成：Lcls 是摩擦等级和风险等级分类损失；Laux 是湿度、雪冰、材质、平整度等辅助任务损失；"
        "Linterval 约束预测区间覆盖弱标签给出的摩擦范围；Lmono 是风险单调性约束；Lcal 是校准相关约束；"
        "λaux、λint、λmono、λcal 是各项损失权重。",
    )
    add_formula(doc, "L_interval = max(0, μ_l - μ_min) + max(0, μ_max - μ_u) + β(μ_max - μ_min)")
    add_para(
        doc,
        "其中 [μl, μu] 是由公开标签映射得到的弱摩擦目标区间；第一项惩罚预测下界过高，第二项惩罚预测上界过低，"
        "第三项惩罚区间过宽。β 是宽度惩罚系数。β 太小会导致区间过宽，β 太大会导致覆盖率下降。",
    )
    add_formula(doc, "I_cal(x) = [ μ_min(x) - q_(1-α),  μ_max(x) + q_(1-α) ]")
    add_para(
        doc,
        "这是校准后的区间。q(1-α) 是在验证集上根据目标覆盖率 1-α 估计得到的 conformal radius。"
        "它的意义是用验证集误差估计一个安全余量，使测试集区间覆盖率更接近目标覆盖率。",
    )

    add_h1(doc, "五、当前核心实验结果")
    fair_rows = []
    for r in report["fair_rows"]:
        fair_rows.append(
            [
                r["dataset"],
                r["faf_status"],
                r["base_status"],
                pct(r["faf_friction"]),
                pct(r["base_friction"]),
                delta(r["faf_friction"], r["base_friction"]),
                pct(r["faf_risk"]),
                pct(r["base_risk"]),
                delta(r["faf_risk"], r["base_risk"]),
                delta(r["faf_cov"], r["base_cov"]),
            ]
        )
    add_table(doc, ["数据集", "FAF状态", "基线状态", "FAF摩擦F1", "基线摩擦F1", "差值", "FAF风险F1", "基线风险F1", "差值", "覆盖差"], fair_rows)
    add_para(
        doc,
        "严格判断：RoadSaW 上当前 FAF 比强 ConvNeXt 基线略好，说明显式物理纹理和多任务风险建模在湿滑/材质场景有价值；"
        "RSCD 上强 ConvNeXt 基线超过当前 FAF，尤其原始区间覆盖率高很多，说明复杂模块存在负迁移。"
        "因此当前不能把完整 FAF 包装成全面优于强基线的最终算法，而应把研究主线转为局部证据、弱监督区间安全性和跨数据集失败机理。",
    )

    add_h1(doc, "六、RoadSC 强基线当前状态")
    roadsc = report["roadsc_state"]
    roadsc_detail = report.get("roadsc_detail") or {}
    roadsc_eval = report.get("roadsc_eval") or {}
    if roadsc_detail:
        add_para(
            doc,
            "RoadSC ConvNeXt 强基线已经完成正式 detailed test。结果显示：强基线在风险任务上达到 100%，"
            "但摩擦/雪类细分仍低于当前 FAF；同时强基线原始区间覆盖率更高，说明 RoadSC 上存在分类性能和区间覆盖之间的取舍。",
        )
        add_table(
            doc,
            ["指标", "RoadSC 强基线 test"],
            [
                ["friction macro-F1", pct(task_metric(roadsc_detail, "friction", "macro_f1"))],
                ["risk macro-F1", pct(task_metric(roadsc_detail, "risk", "macro_f1"))],
                ["snow macro-F1", pct(task_metric(roadsc_detail, "snow", "macro_f1"))],
                ["low-friction recall", pct(low_recall(roadsc_detail))],
                ["raw interval coverage", pct(mu_metric(roadsc_detail, "coverage"))],
                ["interval width", num(mu_metric(roadsc_detail, "width_mean"))],
            ],
        )
    elif roadsc_eval:
        add_para(
            doc,
            "RoadSC ConvNeXt 强基线已经完成初步 evaluate test，detailed/bootstrap/audit 仍在后处理。"
            "以下是初步 evaluate 指标，最终宏 F1 以 detailed_test.json 为准。",
        )
        add_table(
            doc,
            ["指标", "RoadSC 强基线 evaluate"],
            [
                ["friction accuracy", pct(roadsc_eval.get("acc_friction"))],
                ["risk accuracy", pct(roadsc_eval.get("acc_risk"))],
                ["snow accuracy", pct(roadsc_eval.get("acc_snow"))],
                ["raw interval coverage", pct(roadsc_eval.get("mu_interval_coverage"))],
                ["interval width", num(roadsc_eval.get("mu_interval_width"))],
            ],
        )
    elif roadsc:
        val = roadsc.get("val_metrics", {})
        add_para(
            doc,
            f"RoadSC ConvNeXt 强基线正在训练中。当前已完成 epoch {roadsc.get('epoch')}/{roadsc.get('epochs')}，"
            f"stale_epochs={roadsc.get('stale_epochs')}。当前最佳验证 loss 为 {num(roadsc.get('best_metric'))}。"
        )
        add_table(
            doc,
            ["指标", "当前验证值"],
            [
                ["friction accuracy", pct(val.get("acc_friction"))],
                ["risk accuracy", pct(val.get("acc_risk"))],
                ["snow accuracy", pct(val.get("acc_snow"))],
                ["raw interval coverage", pct(val.get("mu_interval_coverage"))],
                ["interval width", num(val.get("mu_interval_width"))],
            ],
        )
        add_para(doc, "目前 RoadSC 基线验证集前期很强，但 epoch 5 已出现回落，说明小数据集上有平台期或过拟合趋势；最终结论需要等 best checkpoint 的正式 test 结果。")
    else:
        add_para(doc, "RoadSC ConvNeXt 强基线尚未写出训练状态。")

    add_h1(doc, "七、模块消融结论")
    p0 = (report["closure"].get("p0_ablation") or []) if isinstance(report["closure"], dict) else []
    rows = []
    for r in p0:
        rows.append([r.get("method", "-"), pct(r.get("friction_f1")), pct(r.get("risk_f1")), pct(r.get("low_friction_recall")), pct(r.get("calibrated_coverage")), pct(r.get("worst_dataset_f1")), r.get("decision", "-")])
    add_table(doc, ["模块", "摩擦F1", "风险F1", "低摩擦召回", "校准覆盖", "最差数据集F1", "判断"], rows)
    add_para(doc, "消融实验说明：PhysicsTexture 是当前最值得保留的有效增量；FrictionSet、DG losses 和完整复杂融合没有形成稳定收益，应快速舍弃或重做。")

    add_h1(doc, "八、跨数据集泛化问题")
    lodo = (report["closure"].get("lodo") or []) if isinstance(report["closure"], dict) else []
    rows = []
    for r in lodo:
        rows.append([r.get("held_out", "-"), pct(r.get("friction_f1")), pct(r.get("risk_f1")), pct(r.get("low_friction_recall")), pct(r.get("calibrated_coverage")), r.get("interpretation", "-")])
    add_table(doc, ["留出数据集", "摩擦F1", "风险F1", "低摩擦召回", "校准覆盖", "解释"], rows)
    add_para(
        doc,
        "LODO 结果整体很差，说明不同公开数据集之间存在强烈域偏移和标签体系偏移。"
        "这反而给论文提供了一个重要问题意识：多数据集混训不一定自然提升泛化，必须显式处理数据集捷径和标签映射差异。",
    )

    add_h1(doc, "九、当前算法创新点")
    add_numbered(
        doc,
        [
            "从普通路面状态分类提升到摩擦可供性区间估计，同时报告等级、风险和不确定区间。",
            "把视觉物理纹理显式加入模型，使模型关注湿滑、反光、粗糙、低纹理等与摩擦相关的图像线索。",
            "采用多任务弱监督，把摩擦、风险、湿度、雪冰、材质、平整度放在同一个学习框架中，让辅助标签解释摩擦来源。",
            "引入区间校准思想，避免只看分类准确率，而同时检查区间覆盖率和区间宽度。",
            "从语义分割、区域证据和 masked consistency 方向设计下一代局部证据模型，目标是让模型回答“哪里导致低摩擦”。",
            "用强 ConvNeXt 基线和跨数据集失败分析约束结论，避免只和弱基线比较造成虚假创新。",
        ],
    )

    add_h1(doc, "十、当前存在的问题")
    add_numbered(
        doc,
        [
            "RSCD 上强 ConvNeXt 基线超过当前 FAF，说明完整复杂融合并不稳定，必须裁剪。",
            "LODO 跨数据集结果严重失败，说明模型容易学习数据集风格、图像裁剪和标签体系差异。",
            "公开数据集没有真实摩擦系数，只能做弱监督摩擦区间，论文表述必须严谨。",
            "RoadSC 基线还在运行，最终同数据集公平对比表需要等 test、bootstrap 和审计文件生成后刷新。",
            "当前还缺少热力图、区域证据可视化、按湿度/材质/雪冰分组的错误分析，这些是后续提升可信度的关键。",
        ],
    )

    add_h1(doc, "十一、下一步工作路线")
    add_numbered(
        doc,
        [
            "先等 RoadSC ConvNeXt 强基线完成，刷新同数据集公平对比表。",
            "保留 PhysicsTexture，舍弃当前不稳定的复杂全量 FAF 组合。",
            "优先快速验证 v23-v25 类候选：region mixture evidence、multi-query region evidence、masked query consistency。",
            "如果候选模型不能超过 RSCD 强基线，则切换论文路线为：强基线 + 可解释局部证据 + 区间校准 + 跨数据集失败诊断。",
            "补充可视化：道路区域证据热力图、低摩擦样本错误案例、区间覆盖-宽度曲线、不同数据集和标签组的混淆矩阵。",
            "论文表述上强调弱监督视觉摩擦可供性，不夸大为实测摩擦系数估计。",
        ],
    )

    add_h1(doc, "十二、可直接汇报的结论")
    add_para(
        doc,
        "目前最稳妥的结论是：这个方向可做，但不能简单依靠多数据集混训或堆复杂模块冲顶会。"
        "已有结果证明 PhysicsTexture 对部分场景有效，RoadSaW 上 FAF 相对强基线有小幅收益；"
        "但 RSCD 上强基线更强，LODO 泛化失败明显。因此下一阶段应迅速转向更像语义分割的局部证据建模，"
        "把创新点集中在可解释低摩擦区域、弱监督区间校准和跨数据集域偏移诊断上。"
    )
    add_h1(doc, "十三、可引用资料")
    add_numbered(
        doc,
        [
            "RSCD 官方页：https://thu-rsxd.com/rscd/ 。要点：friction 包含 dry、wet、water、fresh snow、melted snow、ice；material 和 unevenness 作为辅助属性。",
            "RSCD / Road Surface Image Dataset 发布页：https://data.mendeley.com/datasets/w86hvkrzc5 。要点：图像为车载相机采集后裁剪出的道路表面 patch。",
            "RoadSaW 官方页：https://viscoda.com/index.php/de/downloads-de/roadsaw-dataset-de 。要点：标定相机生成鸟瞰路面 patch，含 3 种表面和 4 级湿度。",
            "RoadSaW CVPRW 2022 论文：RoadSaW: A Large-Scale Dataset for Camera-Based Road Surface and Wetness Estimation。要点：用车载相机和 MARWIS 湿度测量建立路面表面/湿度数据集。",
            "Road Surface Friction Estimation for Winter Conditions Utilising General Visual Features, arXiv:2404.16578 / IEEE ITSC 2024。要点：真实连续摩擦估计通常需要路侧相机与光学摩擦传感器真值；DINOv2 + CNN 的视觉基础特征路线值得借鉴。",
            "Extreme-Road-Image-Dataset：https://github.com/sean-shiyuez/Extreme-Road-Image-Dataset 。要点：与 Zhao et al. 2025 图像+动力学摩擦估计论文相关，提供 6 类极端路面图像，可作为后续低摩擦类别补充数据源。",
        ],
    )


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.18


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_formula(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11.5)
    run.italic = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        shade_cell(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9)
    doc.add_paragraph("")


def shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def status_from(detail: dict[str, Any], run_dir: Path) -> str:
    if detail:
        return "complete"
    if (run_dir / "evaluate_test.json").exists():
        return "eval_only"
    if (run_dir / "best.pt").exists():
        return "trained"
    if (run_dir / "training_state.json").exists():
        return "running"
    return "missing"


def task_metric(report: dict[str, Any], task: str, metric: str) -> Any:
    return ((report.get("tasks") or {}).get(task) or {}).get(metric) if isinstance(report, dict) else None


def low_recall(report: dict[str, Any]) -> Any:
    if not isinstance(report, dict):
        return None
    low = report.get("low_friction_detection") or {}
    if low.get("applicable") is False:
        return None
    return low.get("recall")


def mu_metric(report: dict[str, Any], metric: str) -> Any:
    return (report.get("mu_interval") or {}).get(metric) if isinstance(report, dict) else None


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}
    return data if isinstance(data, dict) else {}


def pct(value: Any) -> str:
    if value is None:
        return "-"
    return f"{float(value) * 100:.2f}%"


def num(value: Any) -> str:
    if value is None:
        return "-"
    return f"{float(value):.4f}"


def delta(a: Any, b: Any) -> str:
    if a is None or b is None:
        return "-"
    return f"{(float(a) - float(b)) * 100:+.2f}%"


if __name__ == "__main__":
    main()
